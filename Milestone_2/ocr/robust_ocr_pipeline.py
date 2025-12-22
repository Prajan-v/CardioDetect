"""Robust Multi-Engine OCR Pipeline for Medical Documents.

Engines:
1. PyMuPDF - Digital PDFs (near 100% accurate)
2. python-docx - Word documents (100% accurate)
3. Tesseract + EasyOCR - Images/Scanned PDFs
4. DeepSeek-VL fallback - For difficult cases
"""

from __future__ import annotations

import logging
import re
import warnings
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import cv2
import numpy as np

warnings.filterwarnings("ignore")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentType(Enum):
    DIGITAL_PDF = "digital_pdf"
    SCANNED_PDF = "scanned_pdf"
    IMAGE = "image"
    WORD_DOC = "word_doc"
    UNKNOWN = "unknown"


class OCREngine(Enum):
    PYMUPDF = "pymupdf"
    PYTHON_DOCX = "python_docx"
    TESSERACT = "tesseract"
    EASYOCR = "easyocr"
    DEEPSEEK_VL = "deepseek_vl"


@dataclass
class ExtractedField:
    name: str
    value: Any
    confidence: float
    raw_text: str = ""
    validated: bool = False


@dataclass
class OCRResult:
    success: bool
    document_type: DocumentType = DocumentType.UNKNOWN
    engine_used: OCREngine = OCREngine.TESSERACT
    fields: Dict[str, ExtractedField] = field(default_factory=dict)
    raw_text: str = ""
    overall_confidence: float = 0.0
    quality_score: str = "unknown"
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    fallback_used: bool = False

    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "document_type": self.document_type.value,
            "engine_used": self.engine_used.value,
            "fields": {n: {"value": f.value, "confidence": f.confidence} for n, f in self.fields.items()},
            "overall_confidence": self.overall_confidence,
            "quality_score": self.quality_score,
        }


# Import engines with fallbacks
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# DeepSeek-VL for difficult cases (optional)
try:
    import torch
    from transformers import AutoModelForCausalLM, AutoProcessor
    DEEPSEEK_VL_AVAILABLE = True
except ImportError:
    DEEPSEEK_VL_AVAILABLE = False


# =============================================================================
# Medical Field Patterns
# =============================================================================

FIELD_PATTERNS = {
    "hemoglobin": [
        r"H[ae]moglobin\s*[:\-–—>→]?\s*([\d.]+)\s*(?:\[.\])?\s*(?:g[/\s]*dl|gm[/\s]*dl)?",
        r"Hb\s*[:\-–—>→]?\s*([\d.]+)",
        r"HEMOGLOBIN\s*[\n\r\s]+([\d.]+)",  # HEMOGLOBIN on one line, value on next
        r"Haemoglobin\s*>\s*([\d.]+)",      # > separator
        r"Hemoglobin[^\d]*([\d]+\.[\d]+)",  # Flexible
        # Value above the Haemoglobin/HEMOGLOBIN label (table layouts)
        r"([\d.]+)[\s\S]{0,80}H[ae]moglobin",
    ],
    "rbc": [
        # R.B.C. / RBC Count, value may be on same or next line
        r"R\.?B\.?C\.?\s*(?:Count)?\s*[:\-–—]?\s*[\n\r\s]*([\d.]+)",
        r"Total\s*R\.?B\.?C\.?\s*(?:Count)?\s*[:\-–—]?\s*[\n\r\s]*([\d.]+)",
        r"RBC\s+COUNT\s*[\n\r\s]*([\d.]+)",  # All caps with newline
        r"RBC\s*Count[^\d]*([\d.]+)",        # Flexible
    ],
    "wbc": [
        # W.B.C. / WBC Count, value may be on same or next line
        r"W\.?B\.?C\.?\s*(?:Count)?\s*[:\-–—]?\s*[\n\r\s]*([\d,]+)",
        r"Total\s*W\.?B\.?C\.?\s*(?:Count)?\s*[:\-–—]?\s*[\n\r\s]*([\d,]+)",
        r"Total\s*WBC\s*Count\s*[:\-–—]?\s*[\n\r\s]*([\d,]+)",
        r"WBC\s+COUNT\s*[\n\r\s]*([\d,]+)",  # All caps with newline
        r"RBC\s+INDICES\s*-\s*WBC\s*[\n\r\s]*([\d,]+)",  # Specific format
    ],
    "platelet": [
        r"Platelet\s*(?:Count)?\s*[:\-]?\s*([\d,]+)",
        r"PLT\s*[:\-]?\s*([\d,]+)",
        r"PLATELET\s*COUNT\s*[\n\r\s]*([\d,]+)",
    ],
    "hematocrit": [
        r"H[ae]matocrit\s*[:\-]?\s*([\d.]+)",
        r"PCV\s*[:\-]?\s*([\d.]+)",
        r"HCT\s*[:\-]?\s*([\d.]+)",
        r"PACKED\s*CELL\s*VOLUME\s*[\n\r\s]*([\d.]+)",
    ],
    "mcv": [
        r"M\.?C\.?V\.?\s*[:\-]?\s*([\d.]+)",
        r"RBC\s*INDICES\s*-\s*MCV\s*[\n\r\s]*([\d.]+)",
    ],
    "mch": [
        r"M\.?C\.?H\.?\s*[:\-]?\s*([\d.]+)",
        r"RBC\s*INDICES\s*-\s*MCH\s*[\n\r\s]*([\d.]+)",
    ],
    "mchc": [
        r"M\.?C\.?H\.?C\.?\s*[:\-]?\s*([\d.]+)",
        r"RBC\s*INDICES\s*-\s*MCHC\s*[\n\r\s]*([\d.]+)",
    ],
    "total_cholesterol": [
        r"Total\s*Cholesterol\s*[:\-]?\s*([\d.]+)",
        r"TC\s*[:\-]?\s*([\d.]+)",
    ],
    "hdl": [r"HDL\s*(?:Cholesterol)?\s*[:\-]?\s*([\d.]+)"],
    "ldl": [r"LDL\s*(?:Cholesterol)?\s*[:\-]?\s*([\d.]+)"],
    "triglycerides": [r"Triglyceride[s]?\s*[:\-]?\s*([\d.]+)"],
    "fasting_glucose": [
        r"(?:Fasting\s*)?Glucose\s*[:\-]?\s*([\d.]+)",
        r"RBS\s*[:\-]?\s*([\d.]+)",
        r"Blood\s*Sugar\s*[:\-]?\s*([\d.]+)",
    ],
    "creatinine": [r"Creatinine\s*[:\-]?\s*([\d.]+)"],
    "urea": [r"(?:Blood\s*)?Urea\s*[:\-]?\s*([\d.]+)"],
    "sgot": [r"SGOT\s*[:\-]?\s*([\d.]+)", r"AST\s*[:\-]?\s*([\d.]+)"],
    "sgpt": [r"SGPT\s*[:\-]?\s*([\d.]+)", r"ALT\s*[:\-]?\s*([\d.]+)"],
    "bilirubin": [r"Bilirubin\s*[:\-]?\s*([\d.]+)"],
    "tsh": [r"TSH\s*[:\-]?\s*([\d.]+)"],
    "pt": [r"(?:Prothrombin\s*Time|PT)\s*[:\-]?\s*([\d.]+)"],
    "inr": [r"INR\s*[:\-]?\s*([\d.]+)"],
    # C-Reactive Protein (CRP) → 267.8 and variants
    "crp": [
        r"C[-\s]*Reactive\s*Protein(?:\s*\(CRP\))?\s*[:\-–—→]?\s*([\d.]+)",
        r"CRP\s*[:\-–—→]?\s*([\d.]+)",
    ],
    "neutrophils": [r"Neutrophil[s]?\s*[:\-]?\s*([\d.]+)"],
    "lymphocytes": [r"Lymphocyte[s]?\s*[:\-]?\s*([\d.]+)"],
    "heart_rate": [r"Heart\s*Rate\s*[:\-]?\s*(\d+)", r"Pulse\s*[:\-]?\s*(\d+)"],
    "age": [r"Age\s*[:\-]?\s*(\d+)"],
    "sex": [r"Sex\s*[:\-]?\s*(Male|Female|M|F)\b", r"Gender\s*[:\-]?\s*(Male|Female|M|F)\b"],
}

VALID_RANGES = {
    "hemoglobin": (5.0, 20.0),
    "rbc": (2.0, 8.0),
    "wbc": (1000, 50000),
    "platelet": (50000, 800000),
    "hematocrit": (15.0, 60.0),
    "mcv": (50.0, 120.0),
    "mch": (15.0, 40.0),
    "mchc": (25.0, 40.0),
    "total_cholesterol": (50, 500),
    "hdl": (10, 150),
    "ldl": (20, 400),
    "triglycerides": (30, 1000),
    "fasting_glucose": (30, 600),
    "creatinine": (0.2, 15.0),
    "urea": (5, 200),
    "sgot": (5, 500),
    "sgpt": (5, 500),
    "bilirubin": (0.1, 20.0),
    "tsh": (0.01, 50.0),
    "pt": (8, 30),
    "inr": (0.5, 5.0),
    "crp": (0, 500),
    "neutrophils": (0, 100),
    "lymphocytes": (0, 100),
    "heart_rate": (30, 250),
    "age": (0, 120),
}


# =============================================================================
# Preprocessing Functions
# =============================================================================

def preprocess_basic(gray: np.ndarray) -> np.ndarray:
    _, th = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
    return th


def preprocess_otsu(gray: np.ndarray) -> np.ndarray:
    _, th = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return th


def preprocess_enhanced(gray: np.ndarray) -> np.ndarray:
    denoised = cv2.fastNlMeansDenoising(gray, None, h=10, templateWindowSize=7, searchWindowSize=21)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(denoised)
    _, th = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return th


def preprocess_aggressive(gray: np.ndarray) -> np.ndarray:
    blurred = cv2.medianBlur(gray, 3)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(blurred)
    _, th = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    kernel = np.ones((2, 2), np.uint8)
    return cv2.dilate(th, kernel, iterations=1)


PREPROCESSORS = {
    "basic": preprocess_basic,
    "otsu": preprocess_otsu,
    "enhanced": preprocess_enhanced,
    "aggressive": preprocess_aggressive,
}


# =============================================================================
# Main Pipeline Class
# =============================================================================

class RobustOCRPipeline:
    """Robust multi-engine OCR pipeline for medical documents."""

    def __init__(
        self,
        use_easyocr: bool = True,
        use_vlm_fallback: bool = False,
        confidence_threshold: float = 70.0,
        verbose: bool = False,
    ):
        self.use_easyocr = use_easyocr and EASYOCR_AVAILABLE
        self.use_vlm_fallback = use_vlm_fallback
        self.confidence_threshold = confidence_threshold
        self.verbose = verbose
        self._easyocr_reader = None
        self._vlm_model = None

    def _log(self, msg: str) -> None:
        if self.verbose:
            logger.info(f"[RobustOCR] {msg}")

    def detect_document_type(self, file_path: Path) -> DocumentType:
        """Detect document type from file extension and content."""
        suffix = file_path.suffix.lower()

        if suffix in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"):
            return DocumentType.IMAGE

        if suffix in (".docx", ".doc"):
            return DocumentType.WORD_DOC

        if suffix == ".pdf":
            if PYMUPDF_AVAILABLE:
                try:
                    doc = fitz.open(str(file_path))
                    text = "".join(doc[i].get_text() for i in range(min(3, len(doc))))
                    doc.close()
                    if len(text.strip()) > 100:
                        return DocumentType.DIGITAL_PDF
                    return DocumentType.SCANNED_PDF
                except Exception:
                    return DocumentType.SCANNED_PDF
            return DocumentType.SCANNED_PDF

        return DocumentType.UNKNOWN

    def extract_from_digital_pdf(self, file_path: Path) -> OCRResult:
        """Extract text from digital PDF using PyMuPDF."""
        result = OCRResult(success=False, document_type=DocumentType.DIGITAL_PDF, engine_used=OCREngine.PYMUPDF)

        if not PYMUPDF_AVAILABLE:
            result.errors.append("PyMuPDF not installed")
            return result

        try:
            doc = fitz.open(str(file_path))
            full_text = "\n".join(doc[i].get_text() for i in range(len(doc)))
            doc.close()

            result.raw_text = full_text
            result.overall_confidence = 99.0
            result.quality_score = "excellent"
            result.success = True
            result.fields = self._extract_fields(full_text)
            self._log(f"Digital PDF: {len(full_text)} chars, {len(result.fields)} fields")
        except Exception as e:
            result.errors.append(str(e))

        return result

    def extract_from_word(self, file_path: Path) -> OCRResult:
        """Extract text from Word document."""
        result = OCRResult(success=False, document_type=DocumentType.WORD_DOC, engine_used=OCREngine.PYTHON_DOCX)

        if not PYTHON_DOCX_AVAILABLE:
            result.errors.append("python-docx not installed")
            return result

        try:
            doc = docx.Document(str(file_path))
            paragraphs = [para.text for para in doc.paragraphs]
            full_text = "\n".join(paragraphs)

            for table in doc.tables:
                for row in table.rows:
                    full_text += "\n" + " | ".join(cell.text for cell in row.cells)

            result.raw_text = full_text
            result.overall_confidence = 100.0
            result.quality_score = "excellent"
            result.success = True
            result.fields = self._extract_fields(full_text)
            self._log(f"Word doc: {len(full_text)} chars, {len(result.fields)} fields")
        except Exception as e:
            result.errors.append(str(e))

        return result

    def _run_tesseract(self, image: np.ndarray, psm: int = 6) -> Tuple[str, float]:
        """Run Tesseract OCR."""
        if not TESSERACT_AVAILABLE:
            return "", 0.0

        config = f"--psm {psm} --oem 3"
        text = pytesseract.image_to_string(image, config=config)
        data = pytesseract.image_to_data(image, config=config, output_type=pytesseract.Output.DICT)
        confs = [float(c) for c in data.get("conf", []) if str(c).lstrip("-").isdigit() and float(c) >= 0]
        avg_conf = sum(confs) / len(confs) if confs else 0.0
        return text, avg_conf

    def _run_easyocr(self, image: np.ndarray) -> Tuple[str, float]:
        """Run EasyOCR as fallback."""
        if not self.use_easyocr:
            return "", 0.0

        if self._easyocr_reader is None:
            self._easyocr_reader = easyocr.Reader(["en"], gpu=False)

        results = self._easyocr_reader.readtext(image)
        texts, confs = [], []
        for _, text, conf in results:
            texts.append(text)
            confs.append(conf * 100)

        return " ".join(texts), (sum(confs) / len(confs) if confs else 0.0)

    def _run_deepseek_vl(self, image: np.ndarray) -> Tuple[str, float]:
        """Run DeepSeek-VL for difficult document OCR."""
        if not self.use_vlm_fallback or not DEEPSEEK_VL_AVAILABLE:
            return "", 0.0

        try:
            from PIL import Image as PILImage

            # Lazy load model
            if self._vlm_model is None:
                self._log("Loading DeepSeek-VL model (this may take a while)...")
                model_name = "deepseek-ai/deepseek-vl-1.3b-chat"  # Smaller version
                self._vlm_processor = AutoProcessor.from_pretrained(model_name, trust_remote_code=True)
                self._vlm_model = AutoModelForCausalLM.from_pretrained(
                    model_name,
                    trust_remote_code=True,
                    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                    device_map="auto" if torch.cuda.is_available() else None,
                )
                self._log("DeepSeek-VL model loaded")

            # Convert numpy array to PIL Image
            if len(image.shape) == 3:
                pil_image = PILImage.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
            else:
                pil_image = PILImage.fromarray(image)

            # Create prompt for medical lab report OCR
            prompt = """Extract all text from this medical lab report image. 
Focus on finding and extracting these values with their numbers:
- Hemoglobin (Hb)
- RBC Count
- WBC Count
- Platelet Count
- Glucose/RBS
- Cholesterol, HDL, LDL
- Any other lab values

Format: FieldName: Value"""

            # Process
            inputs = self._vlm_processor(
                text=prompt,
                images=pil_image,
                return_tensors="pt"
            )

            if torch.cuda.is_available():
                inputs = {k: v.cuda() for k, v in inputs.items()}

            outputs = self._vlm_model.generate(**inputs, max_new_tokens=1024)
            text = self._vlm_processor.decode(outputs[0], skip_special_tokens=True)

            self._log(f"DeepSeek-VL extracted {len(text)} chars")
            return text, 95.0  # High confidence for VLM

        except Exception as e:
            self._log(f"DeepSeek-VL error: {e}")
            return "", 0.0

    def extract_from_image(self, image: np.ndarray) -> OCRResult:
        """Extract text from image using Tesseract + EasyOCR cascade."""
        result = OCRResult(success=False, document_type=DocumentType.IMAGE)

        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image

        best_text, best_conf, best_method = "", 0.0, ""

        # Try each preprocessing + PSM combination
        for prep_name, prep_fn in PREPROCESSORS.items():
            preprocessed = prep_fn(gray)
            for psm in [6, 3, 11]:
                try:
                    text, conf = self._run_tesseract(preprocessed, psm)
                    self._log(f"{prep_name} PSM{psm}: conf={conf:.1f}%")
                    if conf > best_conf:
                        best_text, best_conf, best_method = text, conf, f"{prep_name}+PSM{psm}"
                except Exception as e:
                    self._log(f"{prep_name} PSM{psm}: ERROR {e}")

        result.engine_used = OCREngine.TESSERACT

        # Fallback to EasyOCR if confidence too low
        if best_conf < self.confidence_threshold and self.use_easyocr:
            self._log(f"Low confidence ({best_conf:.1f}%), trying EasyOCR...")
            easy_text, easy_conf = self._run_easyocr(gray)
            if easy_conf > best_conf:
                best_text, best_conf, best_method = easy_text, easy_conf, "easyocr"
                result.engine_used = OCREngine.EASYOCR
                result.fallback_used = True

        # Ultimate fallback to DeepSeek-VL for very difficult cases
        if best_conf < 60.0 and self.use_vlm_fallback and DEEPSEEK_VL_AVAILABLE:
            self._log(f"Very low confidence ({best_conf:.1f}%), trying DeepSeek-VL...")
            vlm_text, vlm_conf = self._run_deepseek_vl(image)
            if vlm_conf > best_conf:
                best_text, best_conf, best_method = vlm_text, vlm_conf, "deepseek_vl"
                result.engine_used = OCREngine.DEEPSEEK_VL
                result.fallback_used = True

        result.raw_text = best_text
        result.overall_confidence = best_conf
        result.quality_score = self._assess_quality(best_conf)
        result.success = best_conf >= 50.0
        result.fields = self._extract_fields(best_text)

        self._log(f"Best: {best_method} @ {best_conf:.1f}%, {len(result.fields)} fields")
        return result

    def extract_from_scanned_pdf(self, file_path: Path) -> OCRResult:
        """Extract text from scanned PDF by converting to images."""
        result = OCRResult(success=False, document_type=DocumentType.SCANNED_PDF)

        if not PDF2IMAGE_AVAILABLE:
            result.errors.append("pdf2image not installed")
            return result

        try:
            images = convert_from_path(str(file_path), dpi=300)
            all_text = []
            total_conf = 0.0

            for i, pil_img in enumerate(images):
                img_array = np.array(pil_img)
                img_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                page_result = self.extract_from_image(img_bgr)
                all_text.append(page_result.raw_text)
                total_conf += page_result.overall_confidence

            result.raw_text = "\n\n--- PAGE BREAK ---\n\n".join(all_text)
            result.overall_confidence = total_conf / len(images) if images else 0.0
            result.quality_score = self._assess_quality(result.overall_confidence)
            result.success = result.overall_confidence >= 50.0
            result.fields = self._extract_fields(result.raw_text)
            result.engine_used = OCREngine.TESSERACT

        except Exception as e:
            result.errors.append(str(e))

        return result

    def extract(self, file_path: Union[str, Path]) -> OCRResult:
        """Main extraction method - auto-detects document type."""
        file_path = Path(file_path)

        if not file_path.exists():
            return OCRResult(success=False, errors=[f"File not found: {file_path}"])

        doc_type = self.detect_document_type(file_path)
        self._log(f"Document type: {doc_type.value}")

        if doc_type == DocumentType.DIGITAL_PDF:
            return self.extract_from_digital_pdf(file_path)
        elif doc_type == DocumentType.WORD_DOC:
            return self.extract_from_word(file_path)
        elif doc_type == DocumentType.SCANNED_PDF:
            return self.extract_from_scanned_pdf(file_path)
        elif doc_type == DocumentType.IMAGE:
            img = cv2.imread(str(file_path), cv2.IMREAD_COLOR)
            if img is None:
                return OCRResult(success=False, errors=[f"Failed to load image: {file_path}"])
            return self.extract_from_image(img)
        else:
            return OCRResult(success=False, errors=[f"Unsupported document type: {doc_type}"])

    def _extract_fields(self, text: str) -> Dict[str, ExtractedField]:
        """Extract medical fields using regex patterns."""
        fields = {}

        for field_name, patterns in FIELD_PATTERNS.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if match:
                    try:
                        raw_val = match.group(1).replace(",", "")

                        if field_name == "sex":
                            value = 1 if raw_val.lower() in ("m", "male") else 0
                        else:
                            value = float(raw_val)

                        # Validate range
                        valid_range = VALID_RANGES.get(field_name)
                        validated = True
                        if valid_range and field_name != "sex":
                            if not (valid_range[0] <= value <= valid_range[1]):
                                validated = False

                        if validated or field_name == "sex":
                            fields[field_name] = ExtractedField(
                                name=field_name,
                                value=value,
                                confidence=80.0,
                                raw_text=match.group(0),
                                validated=validated,
                            )
                            break
                    except (ValueError, AttributeError):
                        continue

        return fields

    def _assess_quality(self, confidence: float) -> str:
        if confidence >= 85:
            return "excellent"
        elif confidence >= 70:
            return "good"
        elif confidence >= 50:
            return "moderate"
        else:
            return "poor"


# =============================================================================
# Convenience Function
# =============================================================================

def extract_medical_fields(file_path: Union[str, Path], verbose: bool = False) -> OCRResult:
    """Extract medical fields from any document type."""
    pipeline = RobustOCRPipeline(verbose=verbose)
    return pipeline.extract(file_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        result = extract_medical_fields(sys.argv[1], verbose=True)
        print("\n=== EXTRACTION RESULT ===")
        print(f"Success: {result.success}")
        print(f"Engine: {result.engine_used.value}")
        print(f"Confidence: {result.overall_confidence:.1f}%")
        print(f"Quality: {result.quality_score}")
        print(f"\n=== EXTRACTED FIELDS ({len(result.fields)}) ===")
        for name, field in result.fields.items():
            print(f"  {name}: {field.value} (conf={field.confidence:.0f}%)")
        if result.errors:
            print(f"\nErrors: {result.errors}")
    else:
        print("Usage: python robust_ocr_pipeline.py <document_path>")
